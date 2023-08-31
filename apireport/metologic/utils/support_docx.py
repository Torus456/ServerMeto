# поддержка формирования документа
import cx_Oracle
import docx
import os
import pandas as pd
from datetime import datetime
from django.conf import settings
from docx.shared import Inches
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def fill_dataframe(path, query, connect, project_args, replace_to=[]):
    """
    Заполнить датафрейм на основе запроса
    """
    SQL_QUERY = ""
    with open(os.path.join(path, query), 'r', encoding="utf-8") as file:
        SQL_QUERY = file.read().replace('\n', ' ')
    for fields in replace_to:
        SQL_QUERY = SQL_QUERY.replace(str(fields[0]), str(fields[1]))
    df = pd.read_sql(SQL_QUERY, con=connect, params=project_args)
    result = df.where((pd.notnull(df)), None)
    return result


def set_color_cell_header(cell, style):
    """
    Устанавливаем цвет ячейки, использую для шапки
    """
    cell = cell
    cell.paragraphs[0].style = style
    cell.paragraphs[0].alignment = 1
    cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="bcbcbc"/>'.format(nsdecls('w'))))


def add_object_value(document, df_vsn_cls):
    """Добавляем значения признаков"""
    values = len(df_vsn_cls)
    if values == 0:
        return
    add_paragraph_before_table(document, "Список значений характеристик класса НМЦ")
    # Значения признаков
    table_vsn = document.add_table(rows=values + 1, cols=3)
    table_vsn.style = 'Table Grid'
    table_vsn.autofit = True
    # Шапка для таблицы значений признаков
    cell = table_vsn.cell(0, 0)
    add_header_table_style(cell, "Наименование характеристики")
    cell = table_vsn.cell(0, 1)
    add_header_table_style(cell, "Значение")
    cell = table_vsn.cell(0, 2)
    add_header_table_style(cell, "Краткое значение")
    set_repeat_table_header(table_vsn.rows[0])
    j = 1
    k = 1
    start_union = 1
    union_name = "Наименование признака"
    for vsn in df_vsn_cls.itertuples():
        cell = table_vsn.cell(j, 0)
        add_cell_table_style(cell, vsn.NAME)
        cell = table_vsn.cell(j, 1)
        add_cell_table_style(cell, "" if vsn.VAL is None else vsn.VAL)
        add_simbol_for_obj_value(cell, vsn)
        cell = table_vsn.cell(j, 2)
        add_cell_table_style(cell, "" if vsn.SVAL is None else vsn.SVAL)
        # объединяем ячейки
        if union_name != vsn.NAME:
            union_name = vsn.NAME
            start_union = j
            k = 1
        else:
            a = table_vsn.cell(start_union, 0)
            b = table_vsn.cell(start_union + k, 0)
            A = a.merge(b)
            A.text = union_name
            add_cell_table_style_for_merge(A, union_name)
            k += 1
        j += 1


def add_simbol_for_obj_value(cell, vsn_value):
    if vsn_value.NUM:
        print(vsn_value.NUM)
        p = cell.paragraphs[0]
        run = p.add_run(str(int(vsn_value.NUM)))
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.superscript = True


def add_description_not_need(document, df_description):
    for vsn in df_description.itertuples():
        p = document.add_paragraph()
        fmt = p.paragraph_format
        fmt.line_spacing = Pt(0)
        fmt.space_after = Pt(0)
        run = p.add_run(str(int(vsn.NUM)))
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9)
        run.font.superscript = True
        run = p.add_run(" - " + vsn.KOMMENT)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9)


def add_description_not_need_marnit(document, df_description):
    for vsn in df_description.itertuples():
        p = document.add_paragraph()
        fmt = p.paragraph_format
        fmt.line_spacing = Pt(0)
        fmt.space_after = Pt(0)
        run = p.add_run(str(int(vsn.NUM)))
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9)
        run.font.superscript = True
        run = p.add_run(" - " + vsn.KOMMENT)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9)


def create_docx_with_tepmplate(data_js):
    """
    Собираем методику с шаблоном
    """
    os.environ["NLS_LANG"] = '.AL32UTF8'
    sql_path = settings.BASE_DIR + "/supp/sqlscript/63"
    docs_path = settings.BASE_DIR + "/supp/word_template/" + "63.docx"
    con = cx_Oracle.connect('CS_ART/CS_ART@192.168.54.17:1521/ORA5.INCON.LO')
    print(data_js)
    result = {}
    project_args = {
        "MLT_ID": data_js.get("project_args").get("mlt_id"),
        "CLF_ID": data_js.get("project_args").get("clf_id"),
        "CLS_ID": data_js.get("project_args").get("cls_id"),
        "CFV_ID": data_js.get("project_args").get("cfv_id"),
        "PRJ_ID": data_js.get("project_args").get("prj_id"),
        "CST_ID": data_js.get("project_args").get("cst_id"),
    }
    project_args_obj = {
        "MLT_ID": data_js.get("project_args").get("mlt_id"),
        "CLF_ID": data_js.get("project_args").get("clf_id"),
        "CLS_ID": data_js.get("project_args").get("cls_id"),
        "CFV_ID": data_js.get("project_args").get("cfv_id"),
        "PRJ_ID": data_js.get("project_args").get("prj_id"),
        "INCLF_ID": data_js.get("project_args").get("inclf_id"),
        "AOBJ_ID": data_js.get("project_args").get("aobj_id"),
        "CST_ID": data_js.get("project_args").get("cst_id"),
    }
    # Классы, признаки, значения, объекты, оквед
    df_cls = fill_dataframe(sql_path, 'cls_for_doc.sql', con, project_args)
    df_dvs = fill_dataframe(sql_path, 'dvs_for_doc.sql', con, project_args)
    df_vsn = fill_dataframe(sql_path, 'vsn_for_doc2.sql', con, project_args_obj)
    df_obj = fill_dataframe(sql_path, 'obj_for_doc.sql', con, project_args_obj)
    df_okved = fill_dataframe(sql_path, 'okved_for_doc.sql', con, project_args_obj)
    print(data_js)
    document = docx.Document(docs_path)
    for index, row in df_cls.iterrows():
        code = "XXX" if row["CODE"] is None else row["CODE"]
        document.add_heading(code + " - " + row["NAME"], row["CLV_LEV"])
        if row["SNAME"]:
            document.add_paragraph().add_run()
            p = document.add_paragraph(
                "Шаблон наименования",
                style="List Bullet 2"
            )
            p = document.add_paragraph(
                row["SNAME"]
            )
        df_obj_cls = df_obj.loc[df_obj["CLS_ID"] == row["CLS_ID"]].head(1)
        if len(df_obj_cls) > 0:
            p = document.add_paragraph(
                "Пример наименования",
                style="List Bullet 2"
            )
            rows = len(df_obj_cls)
            table_obj = document.add_table(rows=rows + 1, cols=2)
            table_obj.style = 'Table Grid'
            table_obj.autofit = False
            # Шапка для таблицы объектов эталона
            cell = table_obj.cell(0, 0)
            cell.width = Inches(5.5)
            cell.text = "Наименование"
            set_color_cell_header(cell, "Normal")
            cell = table_obj.cell(0, 1)
            cell.width = Inches(0.5)
            cell.text = "ЕИ"
            set_color_cell_header(cell, "Normal")
            i = 1
            for ind_obj, obj in df_obj_cls.iterrows():
                cell = table_obj.cell(i, 0)
                cell.text = obj["SNAME"]
                cell = table_obj.cell(i, 1)
                cell.width = Inches(0.5)
                if obj["UMS_CODE"]:
                    cell.text = obj["UMS_CODE"]
                i += 1
            document.add_paragraph().add_run()
            p = document.add_paragraph(
                "Перечень атрибутов класса",
                style="List Bullet 2"
            )
            run = p.add_run()
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.underline = True
            # выбираем признаки для класса
            df_attribute_cls = df_dvs.loc[df_dvs["CLS_ID"] == row["CLS_ID"]]
            rows = len(df_attribute_cls)
            table = document.add_table(rows=rows + 1, cols=3)
            table.style = 'Table Grid'
            table.autofit = True
            # Шапка для таблицы признаков
            cell = table.cell(0, 0)
            cell.text = "Наименование атрибута"
            set_color_cell_header(cell, "Normal")
            cell = table.cell(0, 1)
            cell.text = "Тип атрибута"
            set_color_cell_header(cell, "Normal")
            cell = table.cell(0, 2)
            cell.text = "ЕИ"
            set_color_cell_header(cell, "Normal")
            cnt = 1
            for ind_attr, attr in df_attribute_cls.iterrows():
                cell = table.cell(cnt, 0)
                cell.text = attr["NAME"]
                cell = table.cell(cnt, 1)
                cell.text = attr["VALTYPE"]
                cell = table.cell(cnt, 2)
                if attr["UMS_CODE"]:
                    cell.text = attr["UMS_CODE"]
                cnt += 1
            # Значения признаков
            df_vsn_cls = df_vsn.loc[df_vsn["CLS_ID"] == row["CLS_ID"]]
            values = len(df_vsn_cls)
            if values > 0:
                document.add_paragraph().add_run()
                p = document.add_paragraph(
                    "Список значений атрибутов",
                    style="List Bullet 2"
                )
                run = p.add_run()
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.underline = True
                # Значения признаков
                table_vsn = document.add_table(rows=values + 1, cols=3)
                table_vsn.style = 'Table Grid'
                table_vsn.autofit = True
                # Шапка для таблицы значений признаков
                cell = table_vsn.cell(0, 0)
                cell.text = "Наименование атрибута"
                set_color_cell_header(cell, "Normal")
                cell = table_vsn.cell(0, 1)
                cell.text = "Значение"
                set_color_cell_header(cell, "Normal")
                cell = table_vsn.cell(0, 2)
                cell.text = "Обозначение"
                set_color_cell_header(cell, "Normal")
                j = 1
                k = 1
                start_union = 1
                union_name = "Наименование атрибута"
                for ind_vsn, vsn in df_vsn_cls.iterrows():
                    cell = table_vsn.cell(j, 0)
                    cell.text = vsn["NAME"]
                    cell = table_vsn.cell(j, 1)
                    cell.text = vsn["VALUE"]
                    cell = table_vsn.cell(j, 2)
                    cell.text = "" if vsn["SYMSGN"] is None else vsn["SYMSGN"]
                    # объединяем ячейки
                    if union_name != vsn["NAME"]:
                        union_name = vsn["NAME"]
                        start_union = j
                        k = 1
                    else:
                        a = table_vsn.cell(start_union, 0)
                        b = table_vsn.cell(start_union + k, 0)
                        A = a.merge(b)
                        A.text = union_name
                        add_cell_table_style_for_merge(A, union_name)
                        k += 1
                    j += 1
                document.add_paragraph().add_run()
                if row["UMS_CODE"]:
                    p = document.add_paragraph(
                        "Базовая единица измерения",
                        style="List Bullet 2"
                    )
                    table_umscls = document.add_table(rows=2, cols=3)
                    table_umscls.style = 'Table Grid'
                    table_umscls.autofit = True
                    # Шапка для таблицы значений признаков
                    cell = table_umscls.cell(0, 0)
                    cell.text = "Код ОКЕИ"
                    set_color_cell_header(cell, "Normal")
                    cell = table_umscls.cell(0, 1)
                    cell.text = "Наименование "
                    set_color_cell_header(cell, "Normal")
                    cell = table_umscls.cell(0, 2)
                    cell.text = "Обозначение"
                    set_color_cell_header(cell, "Normal")
                    cell = table_umscls.cell(1, 0)
                    cell.text = row["UMS_ID"]
                    cell = table_umscls.cell(1, 1)
                    cell.text = row["UMS_NAME"]
                    cell = table_umscls.cell(1, 2)
                    cell.text = row["UMS_CODE"]
                df_okved_cls = df_okved.loc[df_okved["CLS_ID"] == row["CLS_ID"]][
                    ["OKVED_CODE", "OKVED_NAME"]
                ].drop_duplicates()
                if len(df_okved_cls) > 0:
                    document.add_paragraph().add_run()
                    p = document.add_paragraph(
                        "Общероссийский классификатор видов экономической деятельности (ОКВЭД 2)",
                        style="List Bullet 2"
                    )
                    table_okved = document.add_table(rows=len(df_okved_cls) + 1, cols=2)
                    table_okved.style = 'Table Grid'
                    table_okved.autofit = True
                    # Шапка для таблицы классов ОКВЕД2
                    cell = table_okved.cell(0, 0)
                    cell.text = "Код класса"
                    set_color_cell_header(cell, "Normal")
                    cell = table_okved.cell(0, 1)
                    cell.text = "Наименование класса"
                    set_color_cell_header(cell, "Normal")
                    r = 1
                    for ind_okv, okv in df_okved_cls.iterrows():
                        cell = table_okved.cell(r, 0)
                        cell.text = okv["OKVED_CODE"]
                        cell = table_okved.cell(r, 1)
                        cell.text = okv["OKVED_NAME"]
                        r += 1
                df_okdp_cls = df_okved.loc[
                    df_okved["CLS_ID"] == row["CLS_ID"]][["OKDP_CODE", "OKDP_NAME"]].drop_duplicates()
                if len(df_okved_cls) > 0:
                    document.add_paragraph().add_run()
                    p = document.add_paragraph(
                        "Общероссийский классификатор продукции по видам экономической деятельности (ОКПД 2)",
                        style="List Bullet 2"
                    )
                    table_okved = document.add_table(rows=len(df_okdp_cls) + 1, cols=2)
                    table_okved.style = 'Table Grid'
                    table_okved.autofit = True
                    # Шапка для таблицы классов ОКДП2
                    cell = table_okved.cell(0, 0)
                    cell.text = "Код класса"
                    set_color_cell_header(cell, "Normal")
                    cell = table_okved.cell(0, 1)
                    cell.text = "Наименование класса"
                    set_color_cell_header(cell, "Normal")
                    r = 1
                    for ind_okv, okv in df_okdp_cls.iterrows():
                        cell = table_okved.cell(r, 0)
                        cell.text = okv["OKDP_CODE"]
                        cell = table_okved.cell(r, 1)
                        cell.text = okv["OKDP_NAME"]
                        r += 1
                df_lot_cls = df_okved.loc[
                    df_okved["CLS_ID"] == row["CLS_ID"]][["LOT_CODE"]].drop_duplicates().dropna().head(1)
                if len(df_lot_cls) > 0:
                    document.add_paragraph().add_run()
                    p = document.add_paragraph(
                        "Код НИИ ЛОТ",
                        style="List Bullet 2"
                    )
                    for ind_okv, okv in df_lot_cls.iterrows():
                        p = document.add_paragraph(okv["LOT_CODE"])
        if row["SNAME"]:
            document.add_page_break()

    for paragraph in document.paragraphs:
        if ':КЛАСС:' in paragraph.text:
            paragraph.text = df_cls["CODE"].iloc[0] + ' - ' + df_cls["NAME"].iloc[0]

    path_file = (settings.BASE_DIR +
                 "/upload/Metodika_" + project_args.get("CLS_ID") + "_" +
                 str(datetime.now().strftime("%Y-_%m-%d-%H_%M_%S")) +
                 ".docx")
    document.save(path_file)
    result["path_file"] = path_file
    result["name"] = df_cls["CODE"].iloc[0]
    return result


def create_docx(data_js):
    """
    Создаем методику по фрагменту
    """
    os.environ["NLS_LANG"] = '.AL32UTF8'
    sql_path = settings.BASE_DIR + "/supp/sqlscript/metologic"
    docs_path = settings.BASE_DIR + "/supp/word_template/" + "Schablon64.docx"
    con = cx_Oracle.connect('CS_ART/CS_ART@192.168.54.17:1521/ORA5.INCON.LO')
    result = {}
    project_args = {
        "MLT_ID": data_js.get("project_args").get("mlt_id"),
        "CLF_ID": data_js.get("project_args").get("clf_id"),
        "CLS_ID": data_js.get("project_args").get("cls_id"),
        "CFV_ID": data_js.get("project_args").get("cfv_id"),
        "PRJ_ID": data_js.get("project_args").get("prj_id")
    }
    project_args_obj = {
        "MLT_ID": data_js.get("project_args").get("mlt_id"),
        "CLF_ID": data_js.get("project_args").get("clf_id"),
        "CLS_ID": data_js.get("project_args").get("cls_id"),
        "CFV_ID": data_js.get("project_args").get("cfv_id"),
        "PRJ_ID": data_js.get("project_args").get("prj_id"),
        "INCLF_ID": data_js.get("project_args").get("inclf_id"),
        "AOBJ_ID": data_js.get("project_args").get("aobj_id"),
    }
    # Классы, признаки, значения, объекты, оквед
    df_cls = fill_dataframe(sql_path, 'cls_for_doc.sql', con, project_args)
    df_dvs = fill_dataframe(sql_path, 'dvs_for_doc.sql', con, project_args)
    df_vsn = fill_dataframe(sql_path, 'vsn_for_doc.sql', con, project_args_obj)
    df_obj = fill_dataframe(sql_path, 'obj_for_doc.sql', con, project_args_obj)
    df_okved = fill_dataframe(sql_path, 'okved_for_doc.sql', con, project_args)
    df_okpd = fill_dataframe(sql_path, 'okpd_for_doc.sql', con, project_args)
    document = docx.Document(docs_path)
    # document.add_heading('Методика', 0)
    # Формируем для создания иерархии
    # for index, row in df_cls.iterrows():
    #     code = "XXX" if row["CODE"] is None else row["CODE"]
    #     p = document.add_paragraph(
    #         code + " - " + row["NAME"],
    #         style="List Bullet"
    #     )
    #     p.paragraph_format.left_indent = Inches((row["CLV_LEV"]-1)/4)
    for index, row in df_cls.iterrows():
        code = "XXX" if row["CODE"] is None else row["CODE"]
        document.add_heading(code + " - " + row["NAME"], row["CLV_LEV"])
        if row["SNAME"]:
            # document.add_paragraph().add_run().add_break()
            p = document.add_paragraph(
                "Шаблон краткого наименования",
                style="List Bullet 2"
            )
            p = document.add_paragraph(
                row["SNAME"]
            )
            p = document.add_paragraph(
                "Шаблон полного наименования",
                style="List Bullet 2"
            )
            p = document.add_paragraph(
                row["FNAME"]
            )
        df_obj_cls = df_obj.loc[df_obj["CLS_ID"] == row["CLS_ID"]]
        if len(df_obj_cls) > 0:
            p = document.add_paragraph(
                "Примеры наименований",
                style="List Bullet 2"
            )
            rows = len(df_obj_cls)
            table_obj = document.add_table(rows=rows + 1, cols=3)
            table_obj.style = 'Table Grid'
            table_obj.autofit = False
            # Шапка для таблицы объектов эталона
            cell = table_obj.cell(0, 0)
            cell.text = "Краткое наименование"
            set_color_cell_header(cell, "Normal")
            cell = table_obj.cell(0, 1)
            cell.width = Inches(3.5)
            cell.text = "Полное наименование"
            set_color_cell_header(cell, "Normal")
            cell = table_obj.cell(0, 2)
            cell.width = Inches(0.5)
            cell.text = "ЕИ"
            set_color_cell_header(cell, "Normal")
            i = 1
            for ind_obj, obj in df_obj_cls.iterrows():
                cell = table_obj.cell(i, 0)
                cell.text = obj["SNAME"]
                cell = table_obj.cell(i, 1)
                cell.text = obj["FNAME"]
                cell = table_obj.cell(i, 2)
                cell.width = Inches(0.5)
                if obj["UMS_CODE"]:
                    cell.text = obj["UMS_CODE"]
                i += 1
            document.add_paragraph().add_run().add_break()
            p = document.add_paragraph(
                "Перечень признаков класса ТМЦ",
                style="List Bullet 2"
            )
            run = p.add_run()
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.underline = True
            # выбираем признаки для класса
            df_attribute_cls = df_dvs.loc[df_dvs["CLS_ID"] == row["CLS_ID"]]
            rows = len(df_attribute_cls)
            table = document.add_table(rows=rows + 1, cols=2)
            table.style = 'Table Grid'
            table.autofit = True
            # Шапка для таблицы признаков
            cell = table.cell(0, 0)
            cell.text = "Наименование признака"
            set_color_cell_header(cell, "Normal")
            cell = table.cell(0, 1)
            cell.text = "Тип признака"
            set_color_cell_header(cell, "Normal")
            cnt = 1
            for ind_attr, attr in df_attribute_cls.iterrows():
                cell = table.cell(cnt, 0)
                cell.text = attr["NAME"]
                cell = table.cell(cnt, 1)
                cell.text = attr["VALTYPE"]
                cnt += 1
            # Значения признаков
            df_vsn_cls = df_vsn.loc[df_vsn["CLS_ID"] == row["CLS_ID"]]
            values = len(df_vsn_cls)
            if values > 0:
                document.add_paragraph().add_run().add_break()
                p = document.add_paragraph(
                    "Список значений",
                    style="List Bullet 2"
                )
                run = p.add_run()
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.underline = True
                # Значения признаков
                table_vsn = document.add_table(rows=values + 1, cols=3)
                table_vsn.style = 'Table Grid'
                table_vsn.autofit = True
                # Шапка для таблицы значений признаков
                cell = table_vsn.cell(0, 0)
                cell.text = "Наименование признака"
                set_color_cell_header(cell, "Normal")
                cell = table_vsn.cell(0, 1)
                cell.text = "Значение"
                set_color_cell_header(cell, "Normal")
                cell = table_vsn.cell(0, 2)
                cell.text = "Обозначение"
                set_color_cell_header(cell, "Normal")
                j = 1
                k = 1
                start_union = 1
                union_name = "Наименование признака"
                for ind_vsn, vsn in df_vsn_cls.iterrows():
                    cell = table_vsn.cell(j, 0)
                    cell.text = vsn["NAME"]
                    cell = table_vsn.cell(j, 1)
                    cell.text = vsn["VALUE"]
                    cell = table_vsn.cell(j, 2)
                    cell.text = "" if vsn["SYMSGN"] is None else vsn["SYMSGN"]
                    # объединяем ячейки
                    if union_name != vsn["NAME"]:
                        union_name = vsn["NAME"]
                        start_union = j
                        k = 1
                    else:
                        a = table_vsn.cell(start_union, 0)
                        b = table_vsn.cell(start_union + k, 0)
                        A = a.merge(b)
                        A.text = union_name
                        add_cell_table_style_for_merge(A, union_name)
                        k += 1
                    j += 1
                document.add_paragraph().add_run().add_break()
                if row["UMS_CODE"]:
                    p = document.add_paragraph(
                        "Базовая единица измерения",
                        style="List Bullet 2"
                    )
                    table_umscls = document.add_table(rows=2, cols=2)
                    table_umscls.style = 'Table Grid'
                    table_umscls.autofit = True
                    # Шапка для таблицы значений признаков
                    cell = table_umscls.cell(0, 0)
                    cell.text = "Наименование "
                    set_color_cell_header(cell, "Normal")
                    cell = table_umscls.cell(0, 1)
                    cell.text = "Обозначение"
                    set_color_cell_header(cell, "Normal")
                    cell = table_umscls.cell(1, 0)
                    cell.text = row["UMS_NAME"]
                    cell = table_umscls.cell(1, 1)
                    cell.text = row["UMS_CODE"]
                df_okved_cls = df_okved.loc[
                    df_okved["CLS_ID"] == row["CLS_ID"]
                ][["OKVED_CODE", "OKVED_NAME"]].drop_duplicates()
                if len(df_okved_cls) > 0:
                    document.add_paragraph().add_run().add_break()
                    p = document.add_paragraph(
                        "Общероссийский классификатор видов экономической деятельности (ОКВЭД2)",
                        style="List Bullet 2"
                    )
                    table_okved = document.add_table(rows=len(df_okved_cls) + 1, cols=2)
                    table_okved.style = 'Table Grid'
                    table_okved.autofit = True
                    # Шапка для таблицы классов ОКВЕД2
                    cell = table_okved.cell(0, 0)
                    cell.text = "Код класса"
                    set_color_cell_header(cell, "Normal")
                    cell = table_okved.cell(0, 1)
                    cell.text = "Наименование класса"
                    set_color_cell_header(cell, "Normal")
                    r = 1
                    for ind_okv, okv in df_okved_cls.iterrows():
                        cell = table_okved.cell(r, 0)
                        cell.text = okv["OKVED_CODE"]
                        cell = table_okved.cell(r, 1)
                        cell.text = okv["OKVED_NAME"]
                        r += 1
                df_okpd_cls = df_okpd.loc[
                    df_okpd["CLS_ID"] == row["CLS_ID"]
                ][["OKVED_CODE", "OKVED_NAME"]].drop_duplicates()
                if len(df_okpd_cls) > 0:
                    print(row["CLS_ID"])
                    print(df_okpd_cls.head())
                    document.add_paragraph().add_run().add_break()
                    p = document.add_paragraph(
                        "Общероссийский классификатор продукции (ОКПД2)",
                        style="List Bullet 2"
                    )
                    table_okpd = document.add_table(rows=len(df_okpd_cls) + 1, cols=2)
                    table_okpd.style = 'Table Grid'
                    table_okpd.autofit = True
                    # Шапка для таблицы классов ОКВЕД2
                    cell = table_okpd.cell(0, 0)
                    cell.text = "Код класса"
                    set_color_cell_header(cell, "Normal")
                    cell = table_okpd.cell(0, 1)
                    cell.text = "Наименование класса"
                    set_color_cell_header(cell, "Normal")
                    r = 1
                    for ind_okp, okp in df_okpd_cls.iterrows():
                        cell = table_okpd.cell(r, 0)
                        cell.text = okp["OKVED_CODE"]
                        cell = table_okpd.cell(r, 1)
                        cell.text = okp["OKVED_NAME"]
                        r += 1
        if row["SNAME"]:
            document.add_page_break()

    path_file = (settings.BASE_DIR +
                 "/upload/Metodika_" +
                 project_args.get("CLS_ID") +
                 "_" +
                 str(datetime.now().strftime("%Y-_%m-%d-%H_%M_%S")) +
                 ".docx")
    document.save(path_file)
    result["path_file"] = path_file
    result["name"] = "Hello"
    return result


def create_docx64(data_js):
    """
    Создаем методику по фрагменту
    """
    os.environ["NLS_LANG"] = '.AL32UTF8'
    sql_path = settings.BASE_DIR + "/supp/sqlscript/metologic"
    docs_path = settings.BASE_DIR + "/supp/word_template/" + "Schablon64.docx"
    con = cx_Oracle.connect('CS_ART/CS_ART@192.168.54.17:1521/ORA5.INCON.LO')
    result = {}
    project_args = {
        "MLT_ID": data_js.get("project_args").get("mlt_id"),
        "CLF_ID": data_js.get("project_args").get("clf_id"),
        "CLS_ID": data_js.get("project_args").get("cls_id"),
        "CFV_ID": data_js.get("project_args").get("cfv_id"),
        "PRJ_ID": data_js.get("project_args").get("prj_id")
    }
    project_args_obj = {
        "MLT_ID": data_js.get("project_args").get("mlt_id"),
        "CLF_ID": data_js.get("project_args").get("clf_id"),
        "CLS_ID": data_js.get("project_args").get("cls_id"),
        "CFV_ID": data_js.get("project_args").get("cfv_id"),
        "PRJ_ID": data_js.get("project_args").get("prj_id"),
        "INCLF_ID": data_js.get("project_args").get("inclf_id"),
        "AOBJ_ID": data_js.get("project_args").get("aobj_id"),
    }
    # Классы, признаки, значения, объекты, оквед
    df_cls = fill_dataframe(sql_path, 'cls_for_doc.sql', con, project_args)
    df_dvs = fill_dataframe(sql_path, 'dvs_for_doc.sql', con, project_args)
    df_vsn = fill_dataframe(sql_path, 'vsn_for_doc.sql', con, project_args_obj)
    df_obj = fill_dataframe(sql_path, 'obj_for_doc.sql', con, project_args_obj)
    document = docx.Document(docs_path)
    # добавляем параграф с оглавлением
    add_table_of_contents(document)
    document.add_page_break()

    for row in df_cls.itertuples():
        code = "XXX" if row.CODE is None else row.CODE
        # descr = "" if row["DESCR"] is None else "(" + row["DESCR"] + ")"
        document.add_heading(code + " - " + row.NAME, row.CLV_LEV)
        # document.add_paragraph().add_run().add_break()
        if row.SNAME:
            add_project_template(document, row)
        # Далее заполняем данные для класса
        df_obj_cls = df_obj.loc[df_obj["CLS_ID"] == row.CLS_ID]
        if len(df_obj_cls) > 0:
            add_project_name(document, df_obj_cls)
            document.add_paragraph().add_run().add_break()
            # выбираем признаки для класса
            df_attribute_cls = df_dvs.loc[df_dvs["CLS_ID"] == row.CLS_ID]
            add_dvs_type_and_name(document, df_attribute_cls)
            # Значения признаков
            df_vsn_cls = df_vsn.loc[df_vsn["CLS_ID"] == row.CLS_ID]
            add_object_value(document, df_vsn_cls)
            document.add_paragraph().add_run().add_break()
            if row.UMS_CODE:
                add_table_of_ums(document, row)
        if row.SNAME:
            document.add_page_break()
    # Меняем в заголовке
    for paragraph in document.paragraphs:
        if ':КЛАСС:' in paragraph.text:
            paragraph.text = "Класс: " + df_cls["CODE"].iloc[0] + ' - ' + df_cls["NAME"].iloc[0]
    text_footer = "Класс: " + df_cls["CODE"].iloc[0] + ' - ' + df_cls["NAME"].iloc[0]
    for section in document.sections:
        footer = section.footer
        if ':КЛАСС:' in footer.tables[0].cell(0, 0).text:
            footer.tables[0].cell(0, 0).text = footer.tables[0].cell(0, 0).text.replace(":КЛАСС:", text_footer)
            footer.tables[0].cell(0, 0).paragraphs[0].alignment = 1
    path_file = (
        settings.BASE_DIR +
        "/upload/Metodika_" +
        project_args.get("CLS_ID") +
        "_" +
        str(datetime.now().strftime("%Y-_%m-%d-%H_%M_%S")) +
        ".docx"
    )
    document.save(path_file)
    result["path_file"] = path_file
    result["name"] = str(df_cls["CODE"].iloc[0] + ' - ' + df_cls["NAME"].iloc[0])
    return result


def create_docx67(data_js):
    """
    Создаем методику по фрагменту
    """
    os.environ["NLS_LANG"] = '.AL32UTF8'
    sql_path = settings.BASE_DIR + "/supp/sqlscript/67"
    docs_path = settings.BASE_DIR + "/supp/word_template/" + "Schablon67.docx"
    con = cx_Oracle.connect('CS_ART/CS_ART@192.168.54.17:1521/ORA5.INCON.LO')
    result = {}
    project_args = {
        "MLT_ID": data_js.get("project_args").get("mlt_id"),
        "CLF_ID": data_js.get("project_args").get("clf_id"),
        "CLS_ID": data_js.get("project_args").get("cls_id"),
        "CFV_ID": data_js.get("project_args").get("cfv_id"),
        "PRJ_ID": data_js.get("project_args").get("prj_id")
    }
    project_args_obj = {
        "MLT_ID": data_js.get("project_args").get("mlt_id"),
        "CLF_ID": data_js.get("project_args").get("clf_id"),
        "CLS_ID": data_js.get("project_args").get("cls_id"),
        "CFV_ID": data_js.get("project_args").get("cfv_id"),
        "PRJ_ID": data_js.get("project_args").get("prj_id"),
        "INCLF_ID": data_js.get("project_args").get("inclf_id"),
        "AOBJ_ID": data_js.get("project_args").get("aobj_id"),
    }
    # Классы, признаки, значения, объекты, оквед
    df_cls = fill_dataframe(sql_path, 'cls_for_doc.sql', con, project_args)
    df_dvs = fill_dataframe(sql_path, 'dvs_for_doc.sql', con, project_args)
    df_vsn = fill_dataframe(sql_path, 'vsn_for_doc.sql', con, project_args_obj)
    df_obj = fill_dataframe(sql_path, 'obj_dop_for_doc.sql', con, project_args_obj)
    df_okpd = fill_dataframe(sql_path, 'okpd_for_doc.sql', con, project_args_obj)
    document = docx.Document(docs_path)
    # добавляем параграф с оглавлением
    add_table_of_contents(document)
    document.add_page_break()

    for row in df_cls.itertuples():
        code = "XXX" if row.CODE is None else row.CODE
        # descr = "" if row["DESCR"] is None else "(" + row["DESCR"] + ")"
        document.add_heading(code + " - " + row.NAME, row.CLV_LEV)
        cls_text = "«" + code + " - " + row.NAME + "»"
        # document.add_paragraph().add_run().add_break()
        if row.SNAME:
            like_heading(document, "1. Схема условного обозначения класса " + cls_text)
            add_project_template(document, row, row.INDICATOR)
        df_obj_cls = df_obj.loc[df_obj["CLS_ID"] == row.CLS_ID]
        # Далее заполняем данные для класса
        if row.ISLEAF == 1:
            like_heading(document, "2. Расшифровка класса " + cls_text)
            # выбираем признаки для класса
            df_attribute_cls = df_dvs.loc[df_dvs["CLS_ID"] == row.CLS_ID]
            add_dvs_type_and_name(document, df_attribute_cls)
            # Значения признаков
            df_vsn_cls = df_vsn.loc[df_vsn["CLS_ID"] == row.CLS_ID]
            add_object_value(document, df_vsn_cls)
            document.add_paragraph().add_run().add_break()
            # Вывод ОКПД2
            df_okpd_cls = df_okpd.loc[df_okpd["CLS_ID"] == row.CLS_ID]
            add_table_of_okpd(document, df_okpd_cls)
            document.add_paragraph().add_run().add_break()
            # Вывод ОКВЭД2
            df_okved_cls = df_okpd.loc[df_okpd["CLS_ID"] == row.CLS_ID]
            add_table_of_okved(document, df_okved_cls)
            document.add_paragraph().add_run().add_break()
            if row.UMS_CODE:
                add_table_of_ums(document, row)
            # df_obj_cls = df_obj.loc[df_obj["CLS_ID"] == row.CLS_ID]
            if len(df_obj_cls) > 0:
                document.add_paragraph().add_run().add_break()
                like_heading(document, "3. Примеры сформированного наименования класса " + cls_text)
                add_object_name_with_value(document, df_obj_cls, row.INDICATOR)
        if row.SNAME:
            document.add_page_break()
    # Меняем в заголовке
    for paragraph in document.paragraphs:
        if ':КЛАСС:' in paragraph.text:
            paragraph.text = "Класс: " + df_cls["CODE"].iloc[0] + ' - ' + df_cls["NAME"].iloc[0]
    text_footer = "Класс: " + df_cls["CODE"].iloc[0] + ' - ' + df_cls["NAME"].iloc[0]
    for section in document.sections:
        footer = section.footer
        if ':КЛАСС:' in footer.tables[0].cell(0, 0).text:
            footer.tables[0].cell(0, 0).text = footer.tables[0].cell(0, 0).text.replace(":КЛАСС:", text_footer)
            footer.tables[0].cell(0, 0).paragraphs[0].alignment = 1
    path_file = (
        settings.BASE_DIR +
        "/upload/Metodika_" +
        project_args.get("CLS_ID") +
        "_" +
        str(datetime.now().strftime("%Y-_%m-%d-%H_%M_%S")) +
        ".docx"
    )
    document.save(path_file)
    result["path_file"] = path_file
    result["name"] = str(df_cls["CODE"].iloc[0] + ' - ' + df_cls["NAME"].iloc[0])
    return result


def insert_table_of_contents(paragraph: object):
    """
    Добавляем оглавление. Создаем оглавление, которое нужно обновить вручную
    """
    paragraph = paragraph
    run = paragraph.add_run()
    fldChar = parse_xml(r'<w:fldChar {} w:fldCharType="begin"/>'.format(nsdecls('w')))  # Добавлем елемент
    instrText = parse_xml(r'<w:instrText {} xml:space="preserve"/>'.format(nsdecls('w')))
    instrText.text = 'TOC \\o "1-9" \\h \\z \\u'
    fldChar2 = parse_xml(r'<w:fldChar {} w:fldCharType="separate"/>'.format(nsdecls('w')))  # creates a new element
    fldChar3 = parse_xml(r'<w:t {} />'.format(nsdecls('w')))
    fldChar3.text = "Правой кнопкой и обновить."
    fldChar2.append(fldChar3)
    fldChar4 = parse_xml(r'<w:fldChar {} w:fldCharType="end"/>'.format(nsdecls('w')))
    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)
    paragraph._p


def add_table_of_okpd(document, row):
    """
    Добавление таблицы с ОКПД
    """
    document.add_paragraph(
        "Общероссийский классификатор видов экономической деятельности (ОКПД2)",
        style="List Bullet 2"
    )
    table_okpdcls = document.add_table(rows=2, cols=2)
    table_okpdcls.style = 'Table Grid'
    table_okpdcls.autofit = True
    # Шапка для таблицы значений признаков
    cell = table_okpdcls.cell(0, 0)
    cell.text = "Код класса"
    set_color_cell_header(cell, "Normal")
    cell = table_okpdcls.cell(0, 1)
    cell.text = "Наименование класса"
    set_color_cell_header(cell, "Normal")
    cell = table_okpdcls.cell(1, 0)
    cell.text = row.COD_OKPD
    cell = table_okpdcls.cell(1, 1)
    cell.text = row.NAME_OKPD


def add_table_of_okved(document, row):
    """
    Добавление таблицы с ОКВЭД
    """
    document.add_paragraph(
        "Общероссийский классификатор видов экономической деятельности (ОКВЭД2)",
        style="List Bullet 2"
    )
    table_okved = document.add_table(rows=2, cols=2)
    table_okved.style = 'Table Grid'
    table_okved.autofit = True
    # Шапка для таблицы значений признаков
    cell = table_okved.cell(0, 0)
    cell.text = "Код класса"
    set_color_cell_header(cell, "Normal")
    cell = table_okved.cell(0, 1)
    cell.text = "Наименование класса"
    set_color_cell_header(cell, "Normal")
    cell = table_okved.cell(1, 0)
    cell.text = row.COD_OKVED
    cell = table_okved.cell(1, 1)
    cell.text = row.NAME_OKVED


def create_docx66(data_js):
    """
    Создаем методику по фрагменту
    """
    os.environ["NLS_LANG"] = '.AL32UTF8'
    sql_path = settings.BASE_DIR + "/supp/sqlscript/66_met"
    docs_path = settings.BASE_DIR + "/supp/word_template/" + "Schablon64.docx"
    con = cx_Oracle.connect('CS_ART/CS_ART@192.168.54.17:1521/ORA5.INCON.LO')
    result = {}
    project_args = {
        "MLT_ID": data_js.get("project_args").get("mlt_id"),
        "CLF_ID": data_js.get("project_args").get("clf_id"),
        "CLS_ID": data_js.get("project_args").get("cls_id"),
        "CFV_ID": data_js.get("project_args").get("cfv_id"),
        "PRJ_ID": data_js.get("project_args").get("prj_id")
    }

    # Классы, признаки, значения, объекты, оквед
    df_cls = fill_dataframe(sql_path, 'cls_for_doc.sql', con, project_args)
    df_dvs = fill_dataframe(sql_path, 'dvs_for_doc.sql', con, project_args)
    df_vsn = fill_dataframe(sql_path, 'vsn_for_doc.sql', con, project_args)
    df_obj = fill_dataframe(sql_path, 'obj_for_doc.sql', con, project_args)
    document = docx.Document(docs_path)
    # добавляем параграф с оглавлением
    add_table_of_contents(document)
    document.add_page_break()

    for index, row in df_cls.iterrows():
        code = "XXX" if row["CODE"] is None else row["CODE"]
        document.add_heading(code + " - " + row["NAME"], row["CLV_LEV"])
        if row["SNAME"]:
            # document.add_paragraph().add_run().add_break()
            p = document.add_paragraph(
                "Шаблон краткого наименования",
                style="List Bullet 2"
            )
            p = document.add_paragraph(
                row["SNAME"]
            )
            p = document.add_paragraph(
                "Шаблон полного наименования",
                style="List Bullet 2"
            )
            p = document.add_paragraph(
                row["FNAME"]
            )
        df_obj_cls = df_obj.loc[df_obj["CLS_ID"] == row["CLS_ID"]]
        if len(df_obj_cls) > 0:
            p = document.add_paragraph(
                "Примеры наименований",
                style="List Bullet 2"
            )
            rows = len(df_obj_cls)
            table_obj = document.add_table(rows=rows + 1, cols=4)
            table_obj.style = 'Table Grid'
            table_obj.autofit = False
            # Шапка для таблицы объектов эталона
            cell = table_obj.cell(0, 0)
            cell.width = Inches(1.0)
            cell.text = "ИД"
            set_color_cell_header(cell, "Normal")
            cell = table_obj.cell(0, 1)
            cell.text = "Краткое наименование"
            set_color_cell_header(cell, "Normal")
            cell = table_obj.cell(0, 2)
            cell.width = Inches(3.5)
            cell.text = "Полное наименование"
            set_color_cell_header(cell, "Normal")
            cell = table_obj.cell(0, 3)
            cell.width = Inches(0.5)
            cell.text = "ЕИ"
            set_color_cell_header(cell, "Normal")
            i = 1
            for ind_obj, obj in df_obj_cls.iterrows():
                cell = table_obj.cell(i, 0)
                cell.width = Inches(1.0)
                cell.text = str(obj["OBJ_ID"])
                cell = table_obj.cell(i, 1)
                cell.text = obj["SNAME"]
                cell = table_obj.cell(i, 2)
                cell.text = obj["FNAME"]
                cell = table_obj.cell(i, 3)
                cell.width = Inches(0.5)
                if obj["UMS_CODE"]:
                    cell.text = obj["UMS_CODE"]
                i += 1
            document.add_paragraph().add_run().add_break()
            p = document.add_paragraph(
                "Перечень признаков класса ТМЦ",
                style="List Bullet 2"
            )
            run = p.add_run()
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.underline = True
            # выбираем признаки для класса
            df_attribute_cls = df_dvs.loc[df_dvs["CLS_ID"] == row["CLS_ID"]]
            rows = len(df_attribute_cls)
            table = document.add_table(rows=rows + 1, cols=2)
            table.style = 'Table Grid'
            table.autofit = True
            # Шапка для таблицы признаков
            cell = table.cell(0, 0)
            cell.text = "Наименование признака"
            set_color_cell_header(cell, "Normal")
            cell = table.cell(0, 1)
            cell.text = "Тип признака"
            set_color_cell_header(cell, "Normal")
            cnt = 1
            for ind_attr, attr in df_attribute_cls.iterrows():
                cell = table.cell(cnt, 0)
                cell.text = attr["NAME"]
                cell = table.cell(cnt, 1)
                cell.text = attr["VALTYPE"]
                cnt += 1
            # Значения признаков
            df_vsn_cls = df_vsn.loc[df_vsn["CLS_ID"] == row["CLS_ID"]]
            values = len(df_vsn_cls)
            if values > 0:
                document.add_paragraph().add_run().add_break()
                p = document.add_paragraph(
                    "Список значений",
                    style="List Bullet 2"
                )
                run = p.add_run()
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.underline = True
                # Значения признаков
                table_vsn = document.add_table(rows=values + 1, cols=3)
                table_vsn.style = 'Table Grid'
                table_vsn.autofit = True
                # Шапка для таблицы значений признаков
                cell = table_vsn.cell(0, 0)
                cell.text = "Наименование признака"
                set_color_cell_header(cell, "Normal")
                cell = table_vsn.cell(0, 1)
                cell.text = "Значение"
                set_color_cell_header(cell, "Normal")
                cell = table_vsn.cell(0, 2)
                cell.text = "Обозначение"
                set_color_cell_header(cell, "Normal")
                j = 1
                k = 1
                start_union = 1
                union_name = "Наименование признака"
                for ind_vsn, vsn in df_vsn_cls.iterrows():
                    cell = table_vsn.cell(j, 0)
                    cell.text = vsn["NAME"]
                    cell = table_vsn.cell(j, 1)
                    cell.text = vsn["VAL"]
                    cell = table_vsn.cell(j, 2)
                    cell.text = "" if vsn["SVAL"] is None else vsn["SVAL"]
                    # объединяем ячейки
                    if union_name != vsn["NAME"]:
                        union_name = vsn["NAME"]
                        start_union = j
                        k = 1
                    else:
                        a = table_vsn.cell(start_union, 0)
                        b = table_vsn.cell(start_union + k, 0)
                        A = a.merge(b)
                        A.text = union_name
                        add_cell_table_style_for_merge(A, union_name)
                        k += 1
                    j += 1
                document.add_paragraph().add_run().add_break()
                if row["UMS_CODE"]:
                    p = document.add_paragraph(
                        "Базовая единица измерения",
                        style="List Bullet 2"
                    )
                    table_umscls = document.add_table(rows=2, cols=2)
                    table_umscls.style = 'Table Grid'
                    table_umscls.autofit = True
                    # Шапка для таблицы значений признаков
                    cell = table_umscls.cell(0, 0)
                    cell.text = "Наименование "
                    set_color_cell_header(cell, "Normal")
                    cell = table_umscls.cell(0, 1)
                    cell.text = "Обозначение"
                    set_color_cell_header(cell, "Normal")
                    cell = table_umscls.cell(1, 0)
                    cell.text = row["UMS_NAME"]
                    cell = table_umscls.cell(1, 1)
                    cell.text = row["UMS_CODE"]
        if row["SNAME"]:
            document.add_page_break()

    # Меняем в заголовке
    for paragraph in document.paragraphs:
        if ':КЛАСС:' in paragraph.text:
            paragraph.text = "Класс: " + df_cls["CODE"].iloc[0] + ' - ' + df_cls["NAME"].iloc[0]
    text_footer = "Класс: " + df_cls["CODE"].iloc[0] + ' - ' + df_cls["NAME"].iloc[0]
    for section in document.sections:
        footer = section.footer
        if ':КЛАСС:' in footer.tables[0].cell(0, 0).text:
            footer.tables[0].cell(0, 0).text = footer.tables[0].cell(0, 0).text.replace(":КЛАСС:", text_footer)
            footer.tables[0].cell(0, 0).paragraphs[0].alignment = 1
    # наименование сохраненного документа
    path_file = (settings.BASE_DIR +
                 "/upload/Metodika_" +
                 project_args.get("CLS_ID") +
                 "_" +
                 str(datetime.now().strftime("%Y-_%m-%d-%H_%M_%S")) +
                 ".docx")

    document.save(path_file)
    result["path_file"] = path_file
    result["name"] = str(df_cls["CODE"].iloc[0] + ' - ' + df_cls["NAME"].iloc[0])
    return result

def create_docx71(data_js):
    """
    Создаем методику по фрагменту
    """
    os.environ["NLS_LANG"] = '.AL32UTF8'
    sql_path = settings.BASE_DIR + "/supp/sqlscript/71"
    docs_path = settings.BASE_DIR + "/supp/word_template/" + "Schablon70.docx"
    con = cx_Oracle.connect('CS_ART/CS_ART@192.168.54.17:1521/ORA5.INCON.LO')
    result = {}
    project_args = {
        "MLT_ID": data_js.get("project_args").get("mlt_id"),
        "CLF_ID": data_js.get("project_args").get("clf_id"),
        "CLS_ID": data_js.get("project_args").get("cls_id"),
        "CFV_ID": data_js.get("project_args").get("cfv_id"),
        "PRJ_ID": data_js.get("project_args").get("prj_id"),
        "INCLF_ID": data_js.get("project_args").get("inclf_id"),
    }
    project_args_obj = {
        "MLT_ID": data_js.get("project_args").get("mlt_id"),
        "CLF_ID": data_js.get("project_args").get("clf_id"),
        "CLS_ID": data_js.get("project_args").get("cls_id"),
        "CFV_ID": data_js.get("project_args").get("cfv_id"),
        "PRJ_ID": data_js.get("project_args").get("prj_id"),
        "INCLF_ID": data_js.get("project_args").get("inclf_id"),
        "AOBJ_ID": data_js.get("project_args").get("aobj_id"),
    }    
    # Классы, признаки, значения, объекты, оквед
    df_cls = fill_dataframe(sql_path, 'cls_for_doc.sql', con, project_args)
    df_dvs = fill_dataframe(sql_path, 'dvs_for_doc.sql', con, project_args)
    df_vsn = fill_dataframe(sql_path, 'vsn_for_doc.sql', con, project_args)
    df_obj = fill_dataframe(sql_path, 'obj_for_doc.sql', con, project_args_obj)
    df_obj_uni = fill_dataframe(sql_path, 'obj_for_doc_uni.sql', con, project_args_obj)
    df_dop = fill_dataframe(sql_path, 'dop_for_doc.sql', con, project_args)
    document = docx.Document(docs_path)
    styles = document.styles
    styles['List Bullet'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    styles['List Bullet'].paragraph_format.space_before = Pt(12)
    styles['Heading 1'].font.name = 'Times New Roman'
    styles['Heading 1'].font.size = Pt(14)
    styles['Heading 1'].font.color.rgb = docx.shared.RGBColor(0, 0, 0)
    styles['Heading 2'].font.name = 'Times New Roman'
    styles['Heading 2'].font.size = Pt(13)
    styles['Heading 2'].font.color.rgb = docx.shared.RGBColor(0, 0, 0)
    styles['Heading 3'].font.name = 'Times New Roman'
    styles['Heading 3'].font.size = Pt(13)
    styles['Heading 3'].font.italic = False
    styles['Heading 3'].font.color.rgb = docx.shared.RGBColor(0, 0, 0)
    styles['Heading 4'].font.name = 'Times New Roman'
    styles['Heading 4'].font.size = Pt(13)
    styles['Heading 4'].font.color.rgb = docx.shared.RGBColor(0, 0, 0)
    styles['Heading 4'].font.italic = False
    styles['Heading 5'].font.name = 'Times New Roman'
    styles['Heading 5'].font.size = Pt(13)
    styles['Heading 5'].font.color.rgb = docx.shared.RGBColor(0, 0, 0)
    styles['Heading 5'].font.italic = False
    # стили уровней
    # добавляем параграф с оглавлением
    # add_table_of_contents(document)
    # document.add_page_break()

    for row in df_cls.itertuples():
        code = "XXX" if row.CODE is None else row.CODE
        document.add_heading(code + " - " + row.NAME, row.CLV_LEV + 1)
        # cls_text = "«" + code + " - " + row.NAME + "»"
        if row.ISLEAF == 1:
            # like_heading(document, "1. Расшифровка класса ")
            # выбираем признаки для класса
            # document.add_paragraph().add_run().add_break()
            df_attribute_cls = df_dvs.loc[df_dvs["CLS_ID"] == row.CLS_ID]
            add_dvs_type_and_name(document, df_attribute_cls)
            # Значения признаков
            df_vsn_cls = df_vsn.loc[df_vsn["CLS_ID"] == row.CLS_ID]
            add_object_value(document, df_vsn_cls)
            df_vsn_not_need = df_vsn.loc[df_vsn["CLS_ID"] == row.CLS_ID][["NUM", "KOMMENT"]].dropna().drop_duplicates()
            add_description_not_need_marnit(document, df_vsn_not_need)
            # Тип атрибута
            df_dop_type_cls = df_dop.loc[df_dop["CLS_ID"] == row.CLS_ID]
            add_dop_values(document, df_dop_type_cls, row)
            # Значения атрибутов
            df_dop_attribute_cls = df_dop.loc[df_dop["CLS_ID"] == row.CLS_ID]
            add_dop_type_and_name(document, df_dop_attribute_cls, row)
            # Шаблоны универсальной записи
            if row.SNAME:
                # document.add_paragraph().add_run().add_break()
                add_project_template_uni(document, row)
            df_obj_uni_cls = df_obj_uni.loc[df_obj_uni["CLS_ID"] == row.CLS_ID]
            # Шаблоны частной записи
            if row.SNAME:
                # document.add_paragraph().add_run().add_break()
                # like_heading(document, "3. Схема условного обозначения класса ")
                add_project_template(document, row)
            df_obj_cls = df_obj.loc[df_obj["CLS_ID"] == row.CLS_ID]
            # Примеры краткого и полного наименований универсальной записи
            if len(df_obj_uni_cls) > 0:
                add_project_name_uni(document, df_obj_uni_cls)
            # Примеры краткого и полного наименований частной записи
            if len(df_obj_cls) > 0:
                add_project_name(document, df_obj_cls)
        if row.SNAME:
            document.add_page_break()
    # Меняем в заголовке
    for paragraph in document.paragraphs:
        if ':КЛАСС:' in paragraph.text:
            paragraph.text = "Категория НМЦ: " + df_cls["CODE"].iloc[0] + ' - ' + df_cls["NAME"].iloc[0]
    for paragraph in document.paragraphs:
        if ':CONTENT:' in paragraph.text:
            paragraph.text = "Оглавление:"
            insert_table_of_contents(paragraph)
            document.add_page_break()
    path_file = (
        settings.BASE_DIR +
        "/upload/Metodika_" +
        project_args.get("CLS_ID") +
        "_" +
        str(datetime.now().strftime("%Y-_%m-%d-%H_%M_%S")) +
        ".docx"
    )
    document.save(path_file)
    result["path_file"] = path_file
    result["name"] = str(df_cls["CODE"].iloc[0] + ' - ' + df_cls["NAME"].iloc[0])
    return result

def create_docx72(data_js):
    """
    Создаем методику по фрагменту
    """
    os.environ["NLS_LANG"] = '.AL32UTF8'
    sql_path = settings.BASE_DIR + "/supp/sqlscript/72"
    docs_path = settings.BASE_DIR + "/supp/word_template/" + "Schablon72.docx"
    con = cx_Oracle.connect('CS_ART/CS_ART@192.168.54.17:1521/ORA5.INCON.LO')
    result = {}
    project_args = {
        "MLT_ID": data_js.get("project_args").get("mlt_id"),
        "CLF_ID": data_js.get("project_args").get("clf_id"),
        "CLS_ID": data_js.get("project_args").get("cls_id"),
        "CFV_ID": data_js.get("project_args").get("cfv_id"),
        "PRJ_ID": data_js.get("project_args").get("prj_id")
    }
    project_args_obj = {
        "MLT_ID": data_js.get("project_args").get("mlt_id"),
        "CLF_ID": data_js.get("project_args").get("clf_id"),
        "CLS_ID": data_js.get("project_args").get("cls_id"),
        "CFV_ID": data_js.get("project_args").get("cfv_id"),
        "PRJ_ID": data_js.get("project_args").get("prj_id"),
        "INCLF_ID": data_js.get("project_args").get("inclf_id"),
    }
    # Классы, признаки, значения, объекты, оквед
    df_cls = fill_dataframe(sql_path, 'cls_for_doc.sql', con, project_args)
    df_dvs = fill_dataframe(sql_path, 'dvs_for_doc.sql', con, project_args)
    df_obj = fill_dataframe(sql_path, 'obj_for_doc.sql', con, project_args_obj)
    df_vsn = fill_dataframe(sql_path, 'vsn_for_doc.sql', con, project_args_obj)
    document = docx.Document(docs_path)
    # добавляем параграф с оглавлением
    add_table_of_contents(document)
    document.add_page_break()

    for row in df_cls.itertuples():
        code = "XXX" if row.CODE is None else row.CODE
        # descr = "" if row["DESCR"] is None else "(" + row["DESCR"] + ")"
        document.add_heading(code + " - " + row.NAME + " - " + row.CUR, row.CLV_LEV)
        cls_text = "«" + code + " - " + row.NAME + " - " + row.CUR + "»"
        # document.add_paragraph().add_run().add_break()
        # Далее заполняем данные для класса
        if row.ISLEAF == 1:
            # выбираем признаки для класса
            df_attribute_cls = df_dvs.loc[df_dvs["CLS_ID"] == row.CLS_ID]
            add_dvs_type_and_name(document, df_attribute_cls)
            # Значения признаков
            df_vsn_cls = df_vsn.loc[df_vsn["CLS_ID"] == row.CLS_ID]
            add_object_value(document, df_vsn_cls)
            document.add_paragraph().add_run().add_break()
        # Шаблон
        if row.SNAME:
            add_project_template(document, row)
        df_obj_cls = df_obj.loc[df_obj["CLS_ID"] == row.CLS_ID]
        # Базовая ЕИ
        if row.UMS_CODE:
            add_table_of_ums(document, row)
        if row.ISLEAF == 1:
            # Проектное sname и ЕИ
            document.add_paragraph().add_run().add_break()
            df_name_cls = df_obj.loc[df_obj["CLS_ID"] == row.CLS_ID]
            add_project_name(document, df_name_cls)
            document.add_paragraph().add_run().add_break()
        if row.ISLEAF == 1:
            add_cls_descr(document, row)
        dd_cls_descr = df_cls.loc[df_cls["CLS_ID"] == row.CLS_ID]
        if row.SNAME:
            document.add_page_break()
    # Меняем в заголовке
    for paragraph in document.paragraphs:
        if ':КЛАСС:' in paragraph.text:
            paragraph.text = '«' + df_cls["CODE"].iloc[0] + ' - ' + df_cls["NAME"].iloc[0] + '»'
    text_footer = '«' + df_cls["CODE"].iloc[0] + ' - ' + df_cls["NAME"].iloc[0] + '»'
    for section in document.sections:
        footer = section.footer
        if ':КЛАСС:' in footer.tables[0].cell(0, 0).text:
            footer.tables[0].cell(0, 0).text = footer.tables[0].cell(0, 0).text.replace(":КЛАСС:", text_footer)
            footer.tables[0].cell(0, 0).paragraphs[0].alignment = 1

    path_file = (
        settings.BASE_DIR +
        "/upload/Metodika_" +
        project_args.get("CLS_ID") +
        "_" +
        str(datetime.now().strftime("%Y-_%m-%d-%H_%M_%S")) +
        ".docx"
    )
    document.save(path_file)
    result["path_file"] = path_file
    result["name"] = str(df_cls["CODE"].iloc[0] + ' - ' + df_cls["NAME"].iloc[0])
    return result


def like_heading(document, header):
    paragr = document.add_paragraph()
    run = paragr.add_run(header)
    run.bold = True


def add_table_of_contents(document: object):
    """
    Добавляем оглавление. Создаем оглавление, которое нужно обновить вручную
    """
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fldChar = parse_xml(r'<w:fldChar {} w:fldCharType="begin"/>'.format(nsdecls('w')))  # Добавлем елемент
    instrText = parse_xml(r'<w:instrText {} xml:space="preserve"/>'.format(nsdecls('w')))
    instrText.text = 'TOC \\o "1-9" \\h \\z \\u'
    fldChar2 = parse_xml(r'<w:fldChar {} w:fldCharType="separate"/>'.format(nsdecls('w')))  # creates a new element
    fldChar3 = parse_xml(r'<w:t {} />'.format(nsdecls('w')))
    fldChar3.text = "Правой кнопкой и обновить."
    fldChar2.append(fldChar3)
    fldChar4 = parse_xml(r'<w:fldChar {} w:fldCharType="end"/>'.format(nsdecls('w')))
    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)
    paragraph._p


def add_table_of_ums(document, row):
    """
    Добавление таблицы с единицей измерения
    """
    document.add_paragraph(
        "Базовая единица измерения",
        style="List Bullet 2"
    )
    table_umscls = document.add_table(rows=2, cols=2)
    table_umscls.style = 'Table Grid'
    table_umscls.autofit = True
    # Шапка для таблицы значений признаков
    cell = table_umscls.cell(0, 0)
    cell.text = "Наименование "
    set_color_cell_header(cell, "Normal")
    cell = table_umscls.cell(0, 1)
    cell.text = "Обозначение"
    set_color_cell_header(cell, "Normal")
    cell = table_umscls.cell(1, 0)
    cell.text = row.UMS_NAME
    cell = table_umscls.cell(1, 1)
    cell.text = row.UMS_CODE


def add_project_template(document, row):
    """
    Добавляем в документ проектные шаблоны наименований
    """
    document.add_paragraph(
        "Шаблон наименования номенклатуры",
        style="List Bullet 2"
    )
    document.add_paragraph(
        row.SNAME
    )


def add_project_name(document, df_obj_cls):
    """
    Добавляем проектные наименования и ЕИ
    """
    df_obj_cls = df_obj_cls.sort_values(by="SNAME")
    document.add_paragraph(
        "Примеры наименований номенклатуры",
        style="List Bullet 2"
    )
    rows = len(df_obj_cls)
    table_obj = document.add_table(rows=rows + 1, cols=2)
    table_obj.style = 'Table Grid'
    table_obj.autofit = False
    # Шапка для таблицы объектов эталона
    cell = table_obj.cell(0, 0)
    cell.text = "Наименование номенклатуры"
    set_color_cell_header(cell, "Normal")
    #cell = table_obj.cell(0, 1)
    cell.width = Inches(5.5)
    #cell.text = "Полное наименование"
    #set_color_cell_header(cell, "Normal")
    cell = table_obj.cell(0, 1)
    cell.width = Inches(0.5)
    cell.text = "ЕИ"
    set_color_cell_header(cell, "Normal")
    i = 1
    for obj in df_obj_cls.itertuples():
        cell = table_obj.cell(i, 0)
        cell.text = obj.SNAME
        #cell = table_obj.cell(i, 1)
        #cell.text = obj.FNAME
        cell = table_obj.cell(i, 1)
        cell.width = Inches(0.5)
        if obj.EI:
            cell.text = obj.EI
        i += 1


def add_object_value(document, df_vsn_cls):
    """Добавляем значения признаков"""
    values = len(df_vsn_cls)
    if values == 0:
        return
    document.add_paragraph().add_run().add_break()
    values = len(df_vsn_cls)
    p = document.add_paragraph(
        "Список значений",
        style="List Bullet 2"
    )
    run = p.add_run()
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.underline = True
    #add_paragraph_before_table(document, "Список значений характеристик класса НМЦ")
    # Значения признаков
    table_vsn = document.add_table(rows=values + 1, cols=2)
    table_vsn.style = 'Table Grid'
    table_vsn.autofit = True
    # Шапка для таблицы значений признаков
    #cell = table_vsn.cell(0, 0)
    #cell.text = "№ П/П"
    #set_color_cell_header(cell, "Normal")
    cell = table_vsn.cell(0, 0)
    cell.text = "Наименование признака"
    set_color_cell_header(cell, "Normal")
    cell = table_vsn.cell(0, 1)
    cell.text = "Значение"
    set_color_cell_header(cell, "Normal")
    #cell = table_vsn.cell(0, 3)
    #cell.text = "Расшифровка"
    #set_color_cell_header(cell, "Normal")
    j = 1
    k = 1
    start_union = 1
    union_name = "Наименование признака"
    for vsn in df_vsn_cls.itertuples():
        #cell = table_vsn.cell(j, 0)
        #cell.text = str(vsn.RN)
        cell = table_vsn.cell(j, 0)
        cell.text = vsn.NAME
        cell = table_vsn.cell(j, 1)
        cell.text = vsn.VAL 
        #cell = table_vsn.cell(j, 3)
        #cell.text = "" if vsn.SYMSGN is None else vsn.SYMSGN
        # объединяем ячейки
        if union_name != vsn.NAME:
            union_name = vsn.NAME
            start_union = j
            k = 1
        else:
            a = table_vsn.cell(start_union, 0)
            b = table_vsn.cell(start_union + k, 0)
            A = a.merge(b)
            A.text = union_name
            k += 1
        j += 1

def add_dvs_type_and_name(document, df_attribute_cls):
    """
    Добавляем в документ Наименования и тип ОД
    """
    p = document.add_paragraph(
        "Перечень признаков класса",
        style="List Bullet 2"
    )
    run = p.add_run()
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.underline = True
    rows = len(df_attribute_cls)
    table = document.add_table(rows=rows + 1, cols=2)
    table.style = 'Table Grid'
    table.autofit = True
    # Шапка для таблицы признаков
    cell = table.cell(0, 0)
    cell.text = "Наименование признака"
    set_color_cell_header(cell, "Normal")
    cell = table.cell(0, 1)
    cell.text = "Тип признака"
    set_color_cell_header(cell, "Normal")
    #cell = table.cell(0, 2)
    #cell.text = "Обязательная"
    #set_color_cell_header(cell, "Normal")
    #cell = table.cell(0, 3)
    #cell.text = "Подчиненная"
    #set_color_cell_header(cell, "Normal")
    cnt = 1
    for attr in df_attribute_cls.itertuples():
        cell = table.cell(cnt, 0)
        cell.text = attr.NAME
        cell = table.cell(cnt, 1)
        cell.text = attr.VALTYPE
        #cell = table.cell(cnt, 2)
        #cell.text = attr.NEED
        #cell = table.cell(cnt, 3)
        #cell.text = attr.DEPEND
        cnt += 1


def add_dop_values(document, df_dop_type_cls, row):
    """
    Добавляем в документ Наименования и тип ОД
    """
    df_dop_type = df_dop_type_cls[["NAME_AT", "TIP"]].drop_duplicates()
    add_paragraph_before_table(document, "Перечень дополнительных атрибутов класса системы классификации НМЦ")
    table = document.add_table(rows=len(df_dop_type) + 2, cols=2)
    table.style = 'Table Grid'
    table.autofit = True
    # Шапка для таблицы признаков
    cell = table.cell(0, 0)
    add_header_table_style(cell, "Наименование атрибута")
    cell = table.cell(0, 1)
    add_header_table_style(cell, "Тип атрибута")
    cnt = 1
    set_repeat_table_header(table.rows[0])
    for attr in df_dop_type.itertuples():
        cell = table.cell(cnt, 0)
        add_cell_table_style(cell, attr.NAME_AT)
        cell = table.cell(cnt, 1)
        add_cell_table_style(cell, attr.TIP)
        cnt += 1
    cell = table.cell(cnt, 0)
    add_cell_table_style(cell, "ЕИ")
    cell = table.cell(cnt, 1)
    add_cell_table_style(cell, "Текстовый")


def add_dop_type_and_name(document, df_dop_attribute_cls, row):
    """
    Добавляем в документ Наименования и тип ОД
    """
    add_paragraph_before_table(document, "Список значений дополнительных атрибутов класса НМЦ")
    rows = len(df_dop_attribute_cls)
    table = document.add_table(rows=rows + 2, cols=2)
    table.style = 'Table Grid'
    table.autofit = True
    # Шапка для таблицы признаков
    cell = table.cell(0, 0)
    add_header_table_style(cell, "Наименование атрибута")
    cell = table.cell(0, 1)
    add_header_table_style(cell, "Значение атрибута")
    cnt = 1
    set_repeat_table_header(table.rows[0])
    j = 1
    k = 1
    start_union = 1
    union_name = "Наименование признака"
    for attr in df_dop_attribute_cls.itertuples():
        cell = table.cell(cnt, 0)
        add_cell_table_style(cell, attr.NAME_AT)
        cell = table.cell(cnt, 1)
        add_cell_table_style(cell, attr.ZNACH_AT)
        cnt += 1
        # объединяем ячейки
        if (union_name != attr.NAME_AT or k == 458):
            union_name = attr.NAME_AT
            start_union = j
            k = 1
        else:
            a = table.cell(start_union, 0)
            b = table.cell(start_union + k, 0)
            A = a.merge(b)
            #add_cell_table_style(A, union_name)
            A.text = union_name
            add_cell_table_style_for_merge(A, union_name)
            k += 1
        j += 1
    cell = table.cell(cnt, 0)
    add_cell_table_style(cell, "ЕИ")
    cell = table.cell(cnt, 1)
    add_cell_table_style(cell, row.UMS_NAME)


def add_project_template_uni(document, row):
    """
    Добавляем в документ проектные шаблоны наименований
    """
    p = document.add_paragraph(
        style="List Bullet"
    )
    run = p.add_run("Шаблон краткого наименования универсальной записи НМЦ")
    run.bold = True
    document.add_paragraph(
        row.SNAME_UNI
    )
    p = document.add_paragraph(
        style="List Bullet"
    )
    run = p.add_run("Шаблон полного наименования универсальной записи НМЦ")
    run.bold = True
    document.add_paragraph(
        row.FNAME_UNI
    )


def add_project_name_uni(document, df_obj_uni_cls):
    """
    Добавляем проектные наименования и ЕИ
    """
    df_obj_uni_cls = df_obj_uni_cls.sort_values(by="SNAME_UNI")
    p = document.add_paragraph(
        style="List Bullet"
    )
    run = p.add_run("Пример наименований универсальной записи НМЦ")
    rows = len(df_obj_uni_cls)
    table_obj = document.add_table(rows=rows + 1, cols=2)
    table_obj.style = 'Table Grid'
    table_obj.autofit = True
    run.bold = True
    # Шапка для таблицы объектов эталона
    cell = table_obj.cell(0, 0)
    cell.width = Inches(3.5)
    add_header_table_style(cell,"Краткое наименование")
    set_color_cell_header(cell, "Normal")
    cell = table_obj.cell(0, 1)
    cell.width = Inches(3.5)
    add_header_table_style(cell, "Полное наименование")
    set_color_cell_header(cell, "Normal")
    i = 1
    set_repeat_table_header(table_obj.rows[0])
    for obj in df_obj_uni_cls.itertuples():
        cell = table_obj.cell(i, 0)
        add_cell_table_style(cell, obj.SNAME_UNI)
        cell = table_obj.cell(i, 1)
        add_cell_table_style(cell, obj.FNAME_UNI)
        i += 1

def add_header_table_style(cell, title):
    """
    Формируем заголовок для таблицы в едином стиле
    """
    p = cell.paragraphs[0]
    p.alignment = 1
    run = p.add_run(title)
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.bold = True
    set_color_cell_header(cell, "Normal")


def add_cell_table_style(cell, title):
    """
    Формируем заголовок для таблицы в едином стиле
    """
    p = cell.paragraphs[0]
    run = p.add_run(title)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)


def set_repeat_table_header(row):
    """ set repeat table row on every new page
    """
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)
    return row

def add_cell_table_style_for_merge(cell, title):
    """
    Формируем заголовок для таблицы в едином стиле
    """
    cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
    cell.paragraphs[0].runs[0].font.size = Pt(11)


def add_paragraph_before_table(document, title):
    p = document.add_paragraph(style="List Bullet")
    run = p.add_run(title)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True


def add_cls_descr(document, row):
    """
    Добавляем в документ
    """
    p = document.add_paragraph(
        style="List Bullet 2"
    )
    run = p.add_run("Спецификация класса")
    #run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    #run.underline = True
    document.add_paragraph(
        row.DESCR
    )